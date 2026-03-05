# -*- coding: utf-8 -*-
"""
将所有 .docx 文件转换为 Markdown 格式，保存到 docs/ 目录
"""
import os
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

BASE = r"C:\Users\lhn\OneDrive\Desktop\江西\弋阳县充换电申报材料\充换电申报材料"
OUT_DIR = r"C:\Users\lhn\OneDrive\Desktop\江西\弋阳充电桩项目\docs"
os.makedirs(OUT_DIR, exist_ok=True)

def docx_to_markdown(filepath):
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        style = para.style.name if para.style else ""
        if "Heading 1" in style or style == "标题 1":
            lines.append(f"# {text}")
        elif "Heading 2" in style or style == "标题 2":
            lines.append(f"## {text}")
        elif "Heading 3" in style or style == "标题 3":
            lines.append(f"### {text}")
        elif "List" in style or "列表" in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)
    # 提取表格
    for i, table in enumerate(doc.tables):
        lines.append(f"\n<!-- 表格 {i+1} -->")
        for j, row in enumerate(table.rows):
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            line = "| " + " | ".join(cells) + " |"
            lines.append(line)
            if j == 0:
                sep = "| " + " | ".join(["---"] * len(cells)) + " |"
                lines.append(sep)
        lines.append("")
    return "\n".join(lines)

docx_files = []
for f in os.listdir(BASE):
    if f.endswith(".docx") and not f.startswith("~$"):
        docx_files.append(os.path.join(BASE, f))

for fp in sorted(docx_files):
    fname = os.path.basename(fp)
    out_name = fname.replace(".docx", ".md")
    out_path = os.path.join(OUT_DIR, out_name)
    try:
        md = docx_to_markdown(fp)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(f"# 原文档：{fname}\n\n")
            f.write(md)
        size = len(md)
        print(f"[OK] {fname} -> {out_name}  ({size} 字符)")
    except Exception as e:
        print(f"[ERR] {fname}: {e}")

print("\n文档转换完成，保存目录:", OUT_DIR)
